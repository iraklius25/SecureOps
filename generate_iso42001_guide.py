"""
Generate: ISO 42001 Implementation Guide using SecureOps Platform
Output: ISO42001_SecureOps_Implementation_Guide.docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Colour palette ──────────────────────────────────────────────────────────
DARK_NAVY   = RGBColor(0x0F, 0x17, 0x2A)
ACCENT_BLUE = RGBColor(0x3B, 0x82, 0xF6)
ACCENT_PURP = RGBColor(0x8B, 0x5C, 0xF6)
ACCENT_TEAL = RGBColor(0x10, 0xB9, 0x81)
ACCENT_AMBE = RGBColor(0xF5, 0x9E, 0x0B)
ACCENT_RED  = RGBColor(0xEF, 0x44, 0x44)
LIGHT_GRAY  = RGBColor(0xF1, 0xF5, 0xF9)
MID_GRAY    = RGBColor(0x64, 0x74, 0x8B)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BLACK       = RGBColor(0x0F, 0x17, 0x2A)

doc = Document()

# ── Page margins ────────────────────────────────────────────────────────────
for section in doc.sections:
    section.page_height    = Cm(29.7)
    section.page_width     = Cm(21.0)
    section.left_margin    = Cm(2.5)
    section.right_margin   = Cm(2.5)
    section.top_margin     = Cm(2.0)
    section.bottom_margin  = Cm(2.0)

# ── Helpers ─────────────────────────────────────────────────────────────────
def set_run_color(run, color): run.font.color.rgb = color
def set_para_shading(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    pPr.append(shd)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color_hex='CCCCCC'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color_hex)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def h1(text, color=DARK_NAVY):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p

def h2(text, color=DARK_NAVY):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

def h3(text, color=ACCENT_BLUE):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    return p

def body(text, bold_parts=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = BLACK
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text, level=0, color=None):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    if color: run.font.color.rgb = color
    p.paragraph_format.left_indent  = Cm(level * 0.6 + 0.5)
    p.paragraph_format.space_after  = Pt(2)
    return p

def numbered(text, level=0):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = BLACK
    p.paragraph_format.left_indent  = Cm(level * 0.6 + 0.5)
    p.paragraph_format.space_after  = Pt(2)
    return p

def note_box(text, bg='EFF6FF', border_color='3B82F6', label='NOTE'):
    p = doc.add_paragraph()
    set_para_shading(p, bg)
    p.paragraph_format.left_indent   = Cm(0.3)
    p.paragraph_format.right_indent  = Cm(0.3)
    p.paragraph_format.space_before  = Pt(6)
    p.paragraph_format.space_after   = Pt(6)
    lbl = p.add_run(f'{label}  ')
    lbl.bold = True
    lbl.font.size = Pt(9)
    lbl.font.color.rgb = ACCENT_BLUE
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = BLACK

def warning_box(text):
    note_box(text, bg='FFF7ED', border_color='F97316', label='⚠  IMPORTANT')

def tip_box(text):
    note_box(text, bg='F0FDF4', border_color='10B981', label='✓  TIP')

def divider(color_hex='3B82F6'):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(8)

def make_table(headers, rows, header_bg='1E293B', col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size  = Pt(9)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row in enumerate(rows):
        trow = t.rows[ri + 1]
        bg   = 'FFFFFF' if ri % 2 == 0 else 'F8FAFC'
        for ci, val in enumerate(row):
            cell = trow.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            if isinstance(val, tuple):
                text, bold = val
                r = p.add_run(text)
                r.bold = bold
                r.font.size = Pt(9.5)
            else:
                r = p.add_run(str(val))
                r.font.size = Pt(9.5)
            r.font.color.rgb = BLACK
            set_cell_borders(cell, 'E2E8F0')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

def phase_banner(number, title, subtitle, color_hex, bg_hex):
    p = doc.add_paragraph()
    set_para_shading(p, bg_hex)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.3)
    r1 = p.add_run(f'PHASE {number}  ')
    r1.bold = True; r1.font.size = Pt(8)
    r1.font.color.rgb = RGBColor.from_string(color_hex)
    r2 = p.add_run(title.upper())
    r2.bold = True; r2.font.size = Pt(14)
    r2.font.color.rgb = DARK_NAVY
    if subtitle:
        p2 = doc.add_paragraph()
        set_para_shading(p2, bg_hex)
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(8)
        p2.paragraph_format.left_indent  = Cm(0.3)
        rs = p2.add_run(subtitle)
        rs.font.size = Pt(10)
        rs.font.color.rgb = MID_GRAY

def step_heading(number, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    nb = p.add_run(f'Step {number} — ')
    nb.bold = True; nb.font.size = Pt(11)
    nb.font.color.rgb = ACCENT_PURP
    nt = p.add_run(title)
    nt.bold = True; nt.font.size = Pt(11)
    nt.font.color.rgb = DARK_NAVY

def platform_ref(where):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f'📍  Platform:  {where}')
    r.italic = True; r.font.size = Pt(9.5)
    r.font.color.rgb = ACCENT_TEAL

# ═══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('SecureOps')
r.bold = True; r.font.size = Pt(36)
r.font.color.rgb = ACCENT_BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('ISO/IEC 42001:2023 Implementation Guide')
r.bold = True; r.font.size = Pt(22)
r.font.color.rgb = DARK_NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('A Step-by-Step Operational Guide Using the SecureOps Platform')
r.font.size = Pt(13); r.font.color.rgb = MID_GRAY

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_shading(p, 'EFF6FF')
r = p.add_run('AI Management System (AIMS) — ISO/IEC 42001:2023')
r.bold = True; r.font.size = Pt(11); r.font.color.rgb = ACCENT_BLUE

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

meta = [
    ('Document Version',  '1.0'),
    ('Standard',          'ISO/IEC 42001:2023'),
    ('Classification',    'Internal — Confidential'),
    ('Prepared by',       'SecureOps Platform'),
    ('Date',              datetime.date.today().strftime('%d %B %Y')),
]
for label, val in meta:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f'{label}:  '); r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = MID_GRAY
    r2 = p.add_run(val); r2.font.size = Pt(10); r2.font.color.rgb = BLACK

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════════════════════
h1('Table of Contents')
divider()
toc = [
    ('1.', 'Introduction & Purpose',                              '3'),
    ('2.', 'About This Guide',                                    '3'),
    ('3.', 'SecureOps Platform — Feature Map to ISO 42001',       '4'),
    ('4.', 'Phase 1 — Foundation & Context (Clause 4)',           '5'),
    ('5.', 'Phase 2 — Leadership & Governance (Clause 5)',        '6'),
    ('6.', 'Phase 3 — Risk Planning (Clause 6)',                  '8'),
    ('7.', 'Phase 4 — Support & Resources (Clause 7)',            '10'),
    ('8.', 'Phase 5 — Operations (Clause 8 / Annex A.9)',         '11'),
    ('9.', 'Phase 6 — Performance Evaluation (Clause 9)',         '13'),
    ('10.','Phase 7 — Continual Improvement (Clause 10)',         '14'),
    ('11.','Phase 8 — Certification Readiness',                   '15'),
    ('12.','12-Month Implementation Timeline',                    '17'),
    ('13.','Certification Body Audit Evidence Checklist',         '18'),
    ('14.','Security & Governance Configuration Checklist',       '19'),
]
for num, title, page in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f'{num}  ')
    r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = ACCENT_BLUE
    r2 = p.add_run(title)
    r2.font.size = Pt(10); r2.font.color.rgb = BLACK
    dots = '.' * max(2, 70 - len(num) - len(title))
    r3 = p.add_run(f'  {dots}  {page}')
    r3.font.size = Pt(10); r3.font.color.rgb = MID_GRAY

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════
h1('1.  Introduction & Purpose')
divider()
body(
    'ISO/IEC 42001:2023 is the first international standard for AI Management Systems (AIMS). '
    'It provides organizations with a systematic framework to responsibly develop, deploy, '
    'monitor, and retire AI systems — addressing risks related to fairness, transparency, '
    'accountability, privacy, and human oversight.'
)
body(
    'This guide walks your organization through a full ISO 42001 implementation using the '
    'SecureOps platform. Every step maps to a specific platform feature, so your team knows '
    'exactly where to go and what to do at every stage of the journey.'
)

h2('Who Is This Guide For?')
bullet('AIMS Programme Managers leading the ISO 42001 implementation')
bullet('AI Ethics Officers responsible for governance and risk oversight')
bullet('IT and Security teams managing AI system infrastructure')
bullet('Internal Auditors conducting AIMS audits')
bullet('Executives and Steering Committee members sponsoring the programme')

h2('What This Guide Covers')
bullet('All 7 clauses of ISO 42001:2023 (Clauses 4–10)')
bullet('All relevant Annex A controls (A.5–A.9)')
bullet('Step-by-step actions in SecureOps for each clause')
bullet('12-month implementation timeline')
bullet('Certification Body (CB) audit evidence checklist')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — ABOUT THIS GUIDE
# ═══════════════════════════════════════════════════════════════════════════════
h1('2.  About This Guide')
divider()
h2('How to Use It')
body('Follow the phases in order. Each phase builds on the last. Do not skip ahead — the output of one phase (e.g. a completed risk register) is the input for the next (e.g. treatment plans).')
note_box('Each step includes a 📍 Platform reference telling you exactly which screen or tab to navigate to in SecureOps.')
h2('Scoring Convention')
body('Use the SecureOps Maturity Assessment (ISO 42001 tab) to score your current state at the start of each phase and again at the end. The target before a certification audit is score ≥ 3 (Defined) across all clauses.')
make_table(
    ['Score', 'Level',       'Meaning'],
    [
        ('0', 'None',        'No evidence of the requirement being addressed'),
        ('1', 'Initial',     'Ad-hoc activity — no documented process'),
        ('2', 'Developing',  'Process exists but not consistently applied'),
        ('3', 'Defined',     'Documented, approved, and consistently followed'),
        ('4', 'Managed',     'Measured and monitored with KPIs'),
        ('5', 'Optimizing',  'Continuously improved based on data'),
    ],
    col_widths=[1.5, 3.0, 10.5],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — PLATFORM FEATURE MAP
# ═══════════════════════════════════════════════════════════════════════════════
h1('3.  SecureOps Platform — Feature Map to ISO 42001')
divider()
body('The table below shows which SecureOps feature covers each ISO 42001 clause and Annex A control.')
make_table(
    ['SecureOps Feature', 'ISO 42001 Clauses / Controls Covered'],
    [
        ('GRC Hub → Document Library',          '4.1, 4.2, 4.3, 5.1, 5.2, 7.5, 8.1, 8.2, 9.2, 9.3, 10.2, A.5, A.7, A.8, A.9'),
        ('GRC Hub → Document Approve button',   '7.5 (documented information control), 9.3 (approved management review minutes)'),
        ('GRC Hub → Document Version History',  '7.5 (version control and retrieval of prior versions)'),
        ('GRC Hub → RACI Matrices',             '5.3, A.6.1, A.6.2 (roles and responsibilities)'),
        ('GRC Hub → Action Tracker (NC fields)','10.2 (root cause, containment, corrective action, verification)'),
        ('GRC Hub → Reviews & Audits',          '9.3 (management review with mandatory agenda checklist)'),
        ('GRC Hub → Supplier Register',         'A.8.1 (third-party AI vendors), supply chain risk'),
        ('GRC Hub → AI Systems Register',       'A.8.1 (AI system inventory), 8.2 (impact assessments), EU AI Act classification'),
        ('Risk Register',                        '6.1.2, 6.1.3, A.5.2 (AI risk identification and treatment)'),
        ('Risk Register → EU AI Act Tier',      'EU AI Act risk classification per AI system or risk'),
        ('Maturity Assessment → ISO 42001',     '6.2 (objectives and KPIs), 9.1 (monitoring and measurement)'),
        ('Settings → LDAP / Active Directory',  'A.6 (access control for AIMS roles)'),
        ('Settings → Audit Log',                '7.5, 9.1 (immutable evidence of who changed what and when)'),
        ('Network Scans / Vulnerabilities',     'A.8.3 (AI system security controls and infrastructure)'),
    ],
    col_widths=[7.0, 9.0],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# PHASE 1
# ═══════════════════════════════════════════════════════════════════════════════
phase_banner(1, 'Foundation & Context', 'Clause 4 — Context of the Organization', '3B82F6', 'EFF6FF')
body('Objective: Define the boundaries of your AIMS before any other work begins. Every subsequent activity depends on having a clear, approved scope.')
divider('3B82F6')

step_heading('1.1', 'Understand your organizational context (Clause 4.1)')
platform_ref('GRC Hub → Document Library → + Add Document')
numbered('Create a new document — Title: "AIMS Context Analysis", Category: Policy')
numbered('Conduct a PESTLE analysis (Political, Economic, Social, Technology, Legal, Environmental) and document:')
bullet('Internal issues: technology maturity, AI skills gap, existing governance structures', 1)
bullet('External issues: regulatory environment (EU AI Act), customer expectations, competitor landscape', 1)
numbered('Set Status to "Draft", assign an Owner and a Review Date (12 months from today)')
numbered('Upload the completed document and use the ✓ Approve button once signed off by the executive sponsor')
tip_box('The Audit Log (Settings → Audit Log) records who approved the document and when — this is the tamper-evident evidence auditors require.')

step_heading('1.2', 'Register all interested parties (Clause 4.2)')
platform_ref('GRC Hub → Document Library → + Add Document')
numbered('Create: Title "AIMS Stakeholder Register", Category: Policy')
numbered('Document each party: Board / executive sponsor, employees, customers, regulators (national AI authority, data protection authority), AI vendors, certification body')
numbered('For each party record: their interest in AI, needs and expectations, how and how often you will engage them')
numbered('Approve the document once complete')

step_heading('1.3', 'Define the AIMS scope (Clause 4.3)')
platform_ref('GRC Hub → Document Library → + Add Document')
numbered('Create: Title "AIMS Scope Statement", Category: Policy')
numbered('Document clearly: which AI systems are in scope (list each by name), which business units and locations are covered, which systems are excluded and why')
warning_box('This document MUST be approved by your executive sponsor before you proceed to Phase 2. An approved, bounded scope is the foundation the certification body audits against.')
numbered('After executive sign-off, click ✓ Approve in the Document Library')

step_heading('1.4', 'Establish your maturity baseline')
platform_ref('Maturity Assessment → ISO 42001 tab')
numbered('Score every item in Clauses 4–10 and all Annex A controls honestly')
numbered('At this stage, most items will score 0–1 — this is expected and correct')
numbered('The radar chart gives you a visual gap analysis; screenshot it as your "Baseline Maturity — [Date]" record')
numbered('Upload the screenshot or export as your baseline evidence document in the Document Library')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# PHASE 2
# ═══════════════════════════════════════════════════════════════════════════════
phase_banner(2, 'Leadership & Governance', 'Clause 5 — Leadership', '8B5CF6', 'F5F3FF')
body('Objective: Secure formal leadership commitment, appoint all mandatory roles, and publish the AI Policy.')
divider('8B5CF6')

step_heading('2.1', 'Appoint mandatory roles and build the RACI matrix (Clause 5.3 / Annex A.6.1)')
platform_ref('GRC Hub → RACI Matrices → + New Matrix')
numbered('Name the matrix: "AIMS Roles and Responsibilities"')
numbered('Add the mandatory AIMS roles as columns:')
bullet('Executive Sponsor (top management)', 1)
bullet('AIMS Programme Manager', 1)
bullet('AI Ethics Officer', 1)
bullet('AIMS Working Group members', 1)
bullet('IT / Data Engineering Lead', 1)
bullet('Legal / Compliance Lead', 1)
bullet('HR Lead', 1)
numbered('Add the key AIMS processes as rows:')
bullet('Define and maintain AIMS scope', 1)
bullet('Maintain AI risk register', 1)
bullet('Conduct AI impact assessments', 1)
bullet('Approve AI system deployments', 1)
bullet('Run internal audits', 1)
bullet('Conduct management reviews', 1)
bullet('Handle AI incidents and nonconformities', 1)
bullet('Manage supplier / vendor AI risk', 1)
numbered('Assign R (Responsible), A (Accountable), C (Consulted), I (Informed) in each cell by clicking the cell to cycle through values')
numbered('Click Export CSV to save the matrix; upload it to the Document Library as "AIMS RACI Matrix"')
note_box('Segregation of duties (Annex A.6.2): ensure the person who develops an AI system is NOT the same person who approves it for deployment. Check your RACI for this conflict.')

step_heading('2.2', 'Write and publish the AI Policy (Clause 5.2 / Annex A.5.1)')
platform_ref('GRC Hub → Document Library → + Add Document')
numbered('Create: Title "AI Policy", Category: Policy')
numbered('The AI Policy must include:')
bullet('Statement of commitment to responsible and ethical AI', 1)
bullet('Eight AI principles: Fairness, Transparency, Accountability, Human oversight, Privacy, Security, Reliability, Sustainability', 1)
bullet('Scope of AI systems and activities covered', 1)
bullet('Reference to relevant regulations (EU AI Act, GDPR Article 22)', 1)
bullet('Top management name and approval signature', 1)
bullet('Annual review commitment', 1)
numbered('Set Status to "Draft" → route for legal review → update to "Review" → get top management sign-off → click ✓ Approve')
numbered('Communicate to ALL in-scope staff — record this in the training register (Step 4.2)')
warning_box('The AI Policy must be approved by the highest level of management. A policy signed only by the AIMS PM is not sufficient for a Stage 2 audit.')

step_heading('2.3', 'Document leadership commitment and establish the Steering Committee (Clause 5.1)')
platform_ref('GRC Hub → Document Library → + Add Document')
numbered('Create: Title "AIMS Governance Charter", Category: Policy')
numbered('Document: executive sponsor name and mandate, Steering Committee composition and terms of reference, meeting frequency (minimum quarterly), resources allocated (AIMS PM FTE, budget envelope)')
numbered('Approve and upload')
numbered('Schedule the first Steering Committee meeting via GRC Hub → Reviews & Audits → + Schedule Review (Type: Management Review)')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# PHASE 3
# ═══════════════════════════════════════════════════════════════════════════════
phase_banner(3, 'Risk Planning', 'Clause 6 — Planning', '10B981', 'F0FDF4')
body('Objective: Identify, assess, and plan the treatment of all AI-related risks. Set measurable AIMS objectives.')
divider('10B981')

step_heading('3.1', 'Build the AI Risk Register (Clause 6.1.2 / Annex A.5.2)')
platform_ref('Risk Register → + Add Risk')
numbered('Add one risk entry for each significant AI-related risk. For every risk, complete:')
make_table(
    ['Field', 'What to Enter', 'Example'],
    [
        ('Title',           'Short descriptive name',                      'Loan model gender bias'),
        ('Category',        'Risk category',                               'Ethical / Compliance'),
        ('Likelihood',      '1–5 scale',                                   '3'),
        ('Impact',          '1–5 scale',                                   '5'),
        ('Treatment',       'Mitigate / Accept / Transfer / Avoid',        'Mitigate'),
        ('EU AI Act Tier',  'Tier of related AI system',                   'High'),
    ],
    col_widths=[3.5, 5.0, 5.5],
)
numbered('SecureOps automatically calculates Risk Score (Likelihood × Impact) and classifies it as Critical / High / Medium / Low')
numbered('Common AI risks to register:')
bullet('Biased model output causing discriminatory decisions (EU AI Act High risk)', 1)
bullet('Model drift degrading accuracy over time without detection', 1)
bullet('Training data poisoning by internal or external actors', 1)
bullet('Prompt injection attacks on LLM-based systems', 1)
bullet('Lack of human oversight on high-stakes automated decisions', 1)
bullet('Third-party AI vendor failure or lock-in', 1)
bullet('Regulatory non-compliance (EU AI Act, GDPR Article 22)', 1)
bullet('Inadequate explainability for regulated decisions (credit, employment, healthcare)', 1)
bullet('Unauthorised AI tool use by staff ("Shadow AI")', 1)

step_heading('3.2', 'Write Risk Treatment Plans (Clause 6.1.3)')
platform_ref('GRC Hub → Action Tracker → + New Task')
numbered('For every Critical and High risk in the register, create an Action Tracker entry:')
bullet('Title: "Treat: [Risk Name]"', 1)
bullet('Source: Select "manual" or "audit" as appropriate', 1)
bullet('NC Type: Select "action"', 1)
bullet('Owner: named individual (not a team)', 1)
bullet('Due Date: within your treatment window (Critical: 30 days, High: 90 days)', 1)
bullet('Click the task row to expand NC fields — document the corrective action planned', 1)
numbered('Update Status as work progresses: Open → In Progress → Completed')
numbered('When treatment is complete, record Verification Evidence in the expanded NC panel')
tip_box('Link treatment actions to the AI System they protect by noting the AI system name in the task description. This traceability is checked by auditors.')

step_heading('3.3', 'Set AIMS Objectives and KPIs (Clause 6.2)')
platform_ref('GRC Hub → Document Library → + Add Document')
numbered('Create: Title "AIMS Objectives and KPIs", Category: Report')
numbered('Define at least 6 SMART objectives. Recommended set:')
make_table(
    ['Ref', 'Objective',                              'KPI Measure',                'Target'],
    [
        ('OBJ-1', 'Achieve ISO 42001 certification',        'Cert audit passed',          'Yes by [date]'),
        ('OBJ-2', 'Maintain complete AI inventory',         '% of AI systems inventoried','100%'),
        ('OBJ-3', 'Regulatory compliance',                  'Open critical findings',     '0'),
        ('OBJ-4', 'Staff competence',                       '% completed AI ethics training','100%'),
        ('OBJ-5', 'Risk reduction',                         'Open Critical/High risks',   'Reduce 50%'),
        ('OBJ-6', 'Ethical AI deployment',                  'Impact assessments done',    '100% pre-deploy'),
    ],
    col_widths=[1.5, 5.5, 5.0, 4.0],
)
numbered('Upload and Approve the document')
numbered('Review KPI performance monthly in the Maturity Assessment dashboard')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# PHASE 4
# ═══════════════════════════════════════════════════════════════════════════════
phase_banner(4, 'Support & Resources', 'Clause 7 — Support', 'F59E0B', 'FFFBEB')
body('Objective: Put in place the people, training, document control, and AI system inventory needed to run the AIMS.')
divider('F59E0B')

step_heading('4.1', 'Build the AI System Inventory (Annex A.8.1)')
platform_ref('GRC Hub → AI Systems tab → + Add AI System')
numbered('Register every AI system in scope. For each system, complete:')
make_table(
    ['Field',              'What to Enter'],
    [
        ('Name',            'Official system name (e.g. "Loan Scoring Model v3")'),
        ('Version',         'Current production version'),
        ('AI Type',         'Generative AI / ML Predictive / NLP / Computer Vision / RPA-AI / Analytics'),
        ('Vendor',          'Developer or vendor name (internal or external)'),
        ('Business Purpose','What decision or task does it support?'),
        ('Decision Role',   'Advisory (human decides) / Automated (AI decides) / Augmented'),
        ('Uses Personal Data','Yes/No — triggers GDPR Article 22 assessment if automated'),
        ('EU AI Act Tier',  'Unacceptable / High / Limited / Minimal'),
        ('Deployment Status','Planned / In Use / Retired / Suspended'),
        ('Owner',           'Named business owner accountable for this system'),
    ],
    col_widths=[4.5, 11.5],
)
numbered('Review the inventory quarterly — mark outdated entries as Retired')
numbered('After completing the impact assessment (Phase 5, Step 5.2), click "Mark Assessed" to record the date and assessor')

step_heading('4.2', 'Document competence and training (Clause 7.2 / 7.3 / Annex A.7.3)')
platform_ref('GRC Hub → Document Library → + Add Document')
numbered('Create: Title "AI Competence Framework" — list required competences by role (developer, business owner, ethics officer, end user, auditor)')
numbered('Create: Title "AI Awareness Training Records" — for each in-scope staff member record: name, role, training completed, date, pass/fail, next renewal date')
numbered('Create actions in GRC Hub → Action Tracker for any outstanding training, with the staff member\'s manager as owner and a due date')
warning_box('100% training completion (OBJ-4) is a hard requirement before a Stage 2 audit. The certification body will spot-check training records for named individuals.')

step_heading('4.3', 'Establish document control (Clause 7.5)')
platform_ref('GRC Hub → Document Library')
body('The Document Library is your AIMS document control system. Apply this naming convention to every document:')
make_table(
    ['Prefix', 'Category',  'Example Title'],
    [
        ('POL',  'Policy',     'POL-001 AI Policy'),
        ('PROC', 'Procedure',  'PROC-001 AI Impact Assessment Procedure'),
        ('REG',  'Register',   'REG-001 AI Risk Register'),
        ('RPT',  'Report',     'RPT-001 Q2 Management Review Minutes'),
        ('REC',  'Record',     'REC-001 Training Completion Records'),
        ('TMPL', 'Template',   'TMPL-001 AI Impact Assessment Template'),
    ],
    col_widths=[2.0, 3.5, 10.5],
)
numbered('Every document must have: a version number, review date, named owner, and Status = Approved before it is used as AIMS evidence')
numbered('When you edit and re-upload a document, SecureOps automatically saves the previous version — auditors can retrieve any historical version if needed')
numbered('Use the ✓ Approve button. The approval is time-stamped and records the approver\'s name — this is the control auditors look for')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# PHASE 5
# ═══════════════════════════════════════════════════════════════════════════════
phase_banner(5, 'Operations', 'Clause 8 — Operation / Annex A.9', 'EF4444', 'FEF2F2')
body('Objective: Control AI systems throughout their full lifecycle. Every AI system must be assessed before deployment.')
divider('EF4444')

step_heading('5.1', 'Write the AI Lifecycle Procedure (Clause 8.1 / Annex A.8.2)')
platform_ref('GRC Hub → Document Library → + Add Document')
numbered('Create: Title "PROC-001 AI Lifecycle Procedure", Category: Procedure')
numbered('The procedure must cover every stage with mandatory controls:')
make_table(
    ['Lifecycle Stage', 'Mandatory Controls'],
    [
        ('Design',         'Requirements review; AI ethics pre-screening; stakeholder consultation'),
        ('Data',           'Data quality assessment; consent and lawful basis verified; data lineage documented'),
        ('Development',    'Bias testing; model card completed; peer review of model logic'),
        ('Testing',        'Performance benchmarks met; fairness metrics within tolerance; adversarial input tests'),
        ('Impact Assessment','Ethics Officer sign-off required before deployment approval (see Step 5.2)'),
        ('Deployment',     'Approval gate: AIMS PM signature on deployment checklist'),
        ('Monitoring',     'Drift detection active; incident alerting configured; model performance reviewed quarterly'),
        ('Retirement',     'Decommission record created; training data deleted per retention policy; system marked Retired in AI Systems register'),
    ],
    col_widths=[4.0, 12.0],
)
numbered('Approve the procedure and upload it to the Document Library')

step_heading('5.2', 'Conduct AI Impact Assessments (Clause 8.2 / Annex A.9.1 / A.9.2)')
platform_ref('GRC Hub → Document Library → + Add Document  |  GRC Hub → AI Systems → Mark Assessed')
numbered('Before deploying any AI system, complete a formal impact assessment. Create one document per system:')
numbered('Title format: "ASSESS-[System Name] Impact Assessment"')
numbered('The assessment must evaluate five dimensions:')
make_table(
    ['Dimension',     'Questions to Answer'],
    [
        ('Individual',  'Could this affect a person\'s rights, access to services, employment, or financial position?'),
        ('Societal',    'Could this reinforce bias at scale or disproportionately affect vulnerable groups?'),
        ('Operational', 'What happens if the system fails, produces incorrect outputs, or becomes unavailable?'),
        ('Legal',       'Does this touch regulated decisions? (credit scoring, employment screening, medical diagnosis, law enforcement)'),
        ('Ethical',     'Is the system transparent, explainable, and fair? Can affected individuals challenge its decisions?'),
    ],
    col_widths=[3.5, 12.5],
)
numbered('Classify the EU AI Act tier (also update the AI Systems register)')
numbered('Obtain sign-off from the AI Ethics Officer')
numbered('Upload and Approve the document')
numbered('Return to GRC Hub → AI Systems, find the system, and click "Mark Assessed" — this records the assessment date and assessor')
warning_box('NO AI system may be deployed to production without a completed and approved impact assessment. Build this gate into your AI Lifecycle Procedure (Step 5.1).')

step_heading('5.3', 'Register and assess AI vendors (Supplier Register)')
platform_ref('GRC Hub → Supplier Register → + Add Supplier')
numbered('For every third-party AI vendor (SaaS AI tools, model providers, data suppliers), create a Supplier record:')
bullet('Risk Rating: assess as Low / Medium / High / Critical based on data access, decision criticality, and dependency', 1)
bullet('DPA: confirm a Data Processing Agreement is in place before any personal data is shared', 1)
bullet('Security Questionnaire: send and receive completed security questionnaire annually', 1)
bullet('Contract End: track renewal dates — set Action Tracker reminders 90 days before expiry', 1)
numbered('Review all High and Critical suppliers annually and update the risk rating')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# PHASE 6
# ═══════════════════════════════════════════════════════════════════════════════
phase_banner(6, 'Performance Evaluation', 'Clause 9 — Performance Evaluation', '6366F1', 'EEF2FF')
body('Objective: Measure whether the AIMS is working and improving. Evidence of monitoring and review is what separates a real AIMS from a paper exercise.')
divider('6366F1')

step_heading('6.1', 'Monitor KPIs monthly (Clause 9.1)')
platform_ref('Maturity Assessment → ISO 42001  |  Risk Register')
numbered('On the first working day of each month:')
bullet('Update Maturity Assessment scores for any clause where work has been completed during the previous month', 1)
bullet('Note the radar chart shape — it should grow outward month on month', 1)
bullet('Open the Risk Register and check: how many Critical/High risks are open vs. last month?', 1)
bullet('Record the KPI values (OBJ-1 to OBJ-6) in your Objectives document', 1)
numbered('Share the radar chart trend with the executive sponsor quarterly')
tip_box('Screenshot the radar chart each month and store the images as evidence in the Document Library under "RPT — Monthly Maturity Snapshot [Month Year]".')

step_heading('6.2', 'Run Internal Audits (Clause 9.2)')
platform_ref('GRC Hub → Document Library  |  GRC Hub → Action Tracker')
numbered('Plan an annual internal audit covering all clauses 4–10 and all Annex A controls. Create:')
bullet('"PROC — Internal Audit Procedure" describing audit methodology, sampling, and reporting', 1)
bullet('"REG — Internal Audit Programme" listing audit subjects, auditors, and scheduled dates', 1)
numbered('For each audit session:')
bullet('Create an audit report document: "RPT — Internal Audit Report [Date]"', 1)
bullet('Record every finding as a Nonconformity (Major NC / Minor NC) or Observation', 1)
numbered('For every Nonconformity, create an Action Tracker entry and fill in the NC fields:')
make_table(
    ['NC Field',               'What to Record'],
    [
        ('NC Type',             'Major NC (systemic failure) / Minor NC (isolated gap) / Observation'),
        ('Source',              'Select "audit"'),
        ('Root Cause',          'WHY did this NC occur? Use 5-Why or fishbone analysis. One sentence minimum.'),
        ('Containment Action',  'Immediate fix to stop the bleeding — done within 24–48 hours'),
        ('Corrective Action',   'Systemic fix to prevent recurrence — addresses the root cause'),
        ('Verification Evidence','How will you prove the corrective action worked? (e.g. re-audit, updated procedure, training record)'),
        ('Verification Date',   'When will you check that the fix has held?'),
    ],
    col_widths=[4.5, 11.5],
)
numbered('Due dates: Major NCs → close within 30 days. Minor NCs → close within 60 days.')
warning_box('Auditors must be independent — they cannot audit areas for which they are responsible. Check your RACI matrix to confirm auditor independence before assigning.')

step_heading('6.3', 'Conduct Management Reviews (Clause 9.3)')
platform_ref('GRC Hub → Reviews & Audits → + Schedule Review')
numbered('Hold at minimum one annual management review (quarterly is strongly recommended)')
numbered('When creating the review in SecureOps, tick all 8 mandatory agenda items in the checklist:')
make_table(
    ['#', 'Mandatory Agenda Item (ISO 42001 Clause 9.3)'],
    [
        ('1', 'Status of actions from previous management reviews'),
        ('2', 'Changes in external and internal context relevant to the AIMS'),
        ('3', 'AI risk register status, trends, and open treatment actions'),
        ('4', 'Results of internal audits (findings, NCs, closure rates)'),
        ('5', 'AIMS objectives and KPI performance (OBJ-1 to OBJ-6)'),
        ('6', 'AI incidents and nonconformities since last review'),
        ('7', 'Opportunities for continual improvement'),
        ('8', 'Adequacy of resources (people, budget, tools)'),
    ],
    col_widths=[1.0, 15.0],
)
numbered('Record decisions and all action items in the Minutes / Notes field')
numbered('Every action from the management review must be created as an Action Tracker entry with a named owner and due date')
numbered('Upload the signed minutes PDF to the Document Library as "RPT — Management Review Minutes [Date]" and click ✓ Approve')
note_box('An unapproved management review record is not acceptable evidence for a Stage 2 audit. The certification body will check that the chair or executive sponsor has formally approved the minutes.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# PHASE 7
# ═══════════════════════════════════════════════════════════════════════════════
phase_banner(7, 'Continual Improvement', 'Clause 10 — Improvement', '0EA5E9', 'F0F9FF')
body('Objective: Embed a PDCA cycle so the AIMS matures continuously, not just at audit time.')
divider('0EA5E9')

step_heading('7.1', 'Maintain an Improvement Register (Clause 10.1)')
platform_ref('GRC Hub → Action Tracker')
numbered('Log every improvement idea using the Action Tracker with NC Type = "Observation". Sources include:')
bullet('Internal audit observations (not just NCs)', 1)
bullet('Management review improvement items', 1)
bullet('Staff suggestions via a named improvement channel', 1)
bullet('Post-AI-incident lessons learned', 1)
bullet('Industry best practice updates (new NIST AI RMF guidance, EU AI Act delegated acts)', 1)
numbered('Prioritise each improvement by impact and effort. Assign to an owner with a realistic due date.')

step_heading('7.2', 'Handle Nonconformities with corrective action (Clause 10.2)')
platform_ref('GRC Hub → Action Tracker → click task row to expand NC fields')
numbered('When a nonconformity is found (at audit, post-incident, or self-detected):')
numbered('Create an Action Tracker entry immediately — do not delay')
numbered('Complete all NC fields within 5 business days:')
bullet('Root cause (use 5-Why — do not accept "human error" as a root cause)', 1)
bullet('Containment action (immediate)', 1)
bullet('Corrective action (systemic)', 1)
bullet('Verification evidence and date', 1)
numbered('Update the Maturity Assessment score for the affected clause once the NC is closed and verified')
numbered('If the same NC recurs, escalate to the Steering Committee — recurring NCs signal a systemic control failure')

step_heading('7.3', 'Run the quarterly improvement cycle')
numbered('Every quarter, run this sequence:')
bullet('Update all Maturity Assessment scores (ISO 42001 tab)', 1)
bullet('Review Risk Register: close treated risks, add newly identified risks', 1)
bullet('Chase overdue Action Tracker items', 1)
bullet('Update KPI values in the Objectives document', 1)
bullet('Present the radar chart trend to the executive sponsor', 1)
bullet('Agree next quarter\'s top 3 improvement priorities', 1)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# PHASE 8
# ═══════════════════════════════════════════════════════════════════════════════
phase_banner(8, 'Certification Readiness', 'Stage 1 & Stage 2 Audit Preparation', 'F97316', 'FFF7ED')
body('Objective: Pass the certification body (CB) Stage 1 (document review) and Stage 2 (implementation audit).')
divider('F97316')

h2('Stage 1 Audit — What the CB Reviews')
body('The CB auditor reviews your documentation to confirm the AIMS is adequately designed before visiting. They will ask for every document listed in Phase 1–7 above. All documents must have Status = Approved.')

h2('Stage 2 Audit — What the CB Tests')
body('The CB auditor interviews staff, observes processes, and samples evidence to confirm the AIMS is actually implemented, not just documented. Common interview questions include:')
bullet('"Show me the last three AI impact assessments."')
bullet('"Walk me through how you identified this risk and what you did to treat it."')
bullet('"Who approved this document and when? How do I know it hasn\'t changed since?"')
bullet('"Show me the training records for your AI development team."')
bullet('"What happened the last time an AI system underperformed? How was it handled?"')

h2('Pre-Audit Self-Check')
note_box('Run this checklist at least 4 weeks before your Stage 1 submission date. Address every gap before submitting.')
make_table(
    ['#', 'Check', 'Pass Criteria', 'Status'],
    [
        ('1',  'All AIMS documents have Status = Approved',           'Zero documents in Draft at audit time',                    '☐'),
        ('2',  'No open Critical/High risks without a treatment plan', 'Every C/H risk has an owner, due date, and corrective action','☐'),
        ('3',  'No overdue Action Tracker items',                     'All items either completed or with an agreed extension date','☐'),
        ('4',  'Internal audit completed in the last 12 months',      'Audit report approved, all NCs closed or on plan',          '☐'),
        ('5',  'Management review conducted in the last 12 months',   'Signed minutes with all 8 agenda items checked',            '☐'),
        ('6',  '100% of in-scope staff have training records',        'Training register current with completion dates',           '☐'),
        ('7',  'AI System Inventory is complete and reviewed',        'Every in-scope AI system registered, EU AI Act tier set',   '☐'),
        ('8',  'Every AI system has a completed impact assessment',   '"Mark Assessed" clicked for every system in AI Systems tab','☐'),
        ('9',  'RACI matrix is approved and published',               'RACI covers all mandatory AIMS processes',                  '☐'),
        ('10', 'Supplier register is current',                        'All High/Critical vendors have active DPA + assessment',    '☐'),
        ('11', 'Maturity Assessment: all clauses score ≥ 3',          'No clause below "Defined" in the radar chart',              '☐'),
        ('12', 'Audit Log reviewed — no unexplained changes',        'Settings → Audit Log checked; no anomalous activity',       '☐'),
    ],
    col_widths=[0.8, 5.5, 6.5, 1.8],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 12 — TIMELINE
# ═══════════════════════════════════════════════════════════════════════════════
h1('12.  12-Month Implementation Timeline')
divider()
make_table(
    ['Month(s)', 'Phase', 'Key Deliverables in SecureOps'],
    [
        ('Month 1',    'Phase 1 — Foundation',      'AIMS Context Analysis, Stakeholder Register, Scope Statement (all Approved); Baseline Maturity snapshot'),
        ('Month 2',    'Phase 2 — Leadership',      'AIMS Governance Charter, AI Policy (Approved); RACI Matrix (all roles assigned); first Steering Committee meeting recorded'),
        ('Month 3–4',  'Phase 3 — Risk Planning',   'Full AI Risk Register; treatment plans as Action Tracker entries; AIMS Objectives & KPIs document'),
        ('Month 5–6',  'Phase 4 — Support',         'AI System Inventory complete; AI Lifecycle Procedure; Training Records; Supplier Register populated'),
        ('Month 7–8',  'Phase 5 — Operations',      'Impact Assessments for all in-scope AI systems (all marked Assessed); Supplier assessments for High/Critical vendors'),
        ('Month 9',    'Phase 6 — Evaluation',      'First Internal Audit completed; all NCs closed; Management Review with signed minutes'),
        ('Month 10',   'Phase 7 — Improvement',     'NC corrective actions verified; Improvement register active; Maturity re-assessed (target: all clauses ≥ 3)'),
        ('Month 11',   'Pre-Certification',         'Pre-audit self-check; all gaps remediated; Stage 1 submission package prepared'),
        ('Month 12',   'Phase 8 — Certification',   'Stage 1 audit (CB document review); Stage 2 audit (CB implementation audit); Certificate issued'),
    ],
    col_widths=[2.5, 4.0, 9.5],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 13 — CB AUDIT EVIDENCE CHECKLIST
# ═══════════════════════════════════════════════════════════════════════════════
h1('13.  Certification Body Audit Evidence Checklist')
divider()
body('This table maps every ISO 42001 clause to the evidence the CB will ask for and where it lives in SecureOps.')
make_table(
    ['Clause', 'Evidence Required', 'SecureOps Location'],
    [
        ('4.1',  'AIMS Context Analysis (PESTLE)',                   'GRC Hub → Document Library'),
        ('4.2',  'Stakeholder Register',                             'GRC Hub → Document Library'),
        ('4.3',  'AIMS Scope Statement (approved)',                  'GRC Hub → Document Library'),
        ('5.1',  'Governance Charter; Steering Committee minutes',   'GRC Hub → Documents + Reviews'),
        ('5.2',  'AI Policy (top management signed)',                'GRC Hub → Document Library'),
        ('5.3',  'RACI Matrix with all roles',                      'GRC Hub → RACI Matrices'),
        ('6.1.2','AI Risk Register with scores',                    'Risk Register'),
        ('6.1.3','Risk treatment plans with owners and due dates',   'GRC Hub → Action Tracker'),
        ('6.2',  'AIMS Objectives and KPIs with evidence',          'GRC Hub → Document Library'),
        ('7.2',  'Competence Framework; training records',           'GRC Hub → Document Library'),
        ('7.3',  'Awareness training records (100% completion)',     'GRC Hub → Document Library'),
        ('7.5',  'Document version control and approval log',        'GRC Hub → Document Library + Settings → Audit Log'),
        ('8.1',  'AI Lifecycle Procedure',                          'GRC Hub → Document Library'),
        ('8.2',  'Impact Assessments per AI system (all approved)',  'GRC Hub → Documents + AI Systems → Mark Assessed'),
        ('9.1',  'KPI monitoring records; maturity snapshots',      'Maturity Assessment + Risk Register'),
        ('9.2',  'Internal Audit report; NC log with closure proof', 'GRC Hub → Documents + Action Tracker'),
        ('9.3',  'Management Review minutes (all 8 items checked)',  'GRC Hub → Reviews & Audits'),
        ('10.2', 'NC log with root cause, corrective action, verification','GRC Hub → Action Tracker (expanded NC fields)'),
        ('A.5.1','AI Policy + AI Ethics Policy',                    'GRC Hub → Document Library'),
        ('A.5.2','AI Risk Management Policy / methodology',         'GRC Hub → Document Library'),
        ('A.6.1','RACI Matrix with segregation of duties evidence',  'GRC Hub → RACI Matrices'),
        ('A.6.2','Evidence that developers ≠ deployment approvers', 'GRC Hub → RACI + Document approvals'),
        ('A.7.3','Role-specific AI training records',               'GRC Hub → Document Library'),
        ('A.8.1','AI System Inventory (all systems, all tiers)',     'GRC Hub → AI Systems tab'),
        ('A.8.2','Acceptable use guidelines; Shadow AI controls',    'GRC Hub → Document Library'),
        ('A.9.1','Impact Assessment methodology / procedure',        'GRC Hub → Document Library'),
        ('A.9.2','Per-system impact assessment records (approved)',   'GRC Hub → Document Library + AI Systems'),
    ],
    col_widths=[1.8, 7.0, 7.2],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 14 — SECURITY & GOVERNANCE CONFIG CHECKLIST
# ═══════════════════════════════════════════════════════════════════════════════
h1('14.  Security & Governance Configuration Checklist')
divider()
body('Before going live with the platform in a production environment, verify these SecureOps settings:')
make_table(
    ['#', 'Configuration Item',                              'Where to Configure',                        'Done'],
    [
        ('1',  'Changed default admin password',              'Settings → User Management',                '☐'),
        ('2',  'JWT_SECRET set to random 64-char hex string', 'Server .env file',                          '☐'),
        ('3',  'Firewall: ports 4000 and 5432 blocked externally','Server: sudo ufw status',               '☐'),
        ('4',  'HTTPS / TLS configured',                      'Settings (INSTALL.md Step 14)',             '☐'),
        ('5',  'LDAP / AD authentication configured',         'Settings → LDAP / Active Directory',       '☐'),
        ('6',  'LDAP group-to-role mappings set correctly',   'Settings → LDAP → Group Mapping table',    '☐'),
        ('7',  'Audit Log reviewed weekly',                   'Settings → Audit Log',                     '☐'),
        ('8',  'Nightly database backup configured',          'Server cron (INSTALL.md → Daily Backup)',   '☐'),
        ('9',  'PM2 startup configured (survives reboots)',   'pm2 startup (INSTALL.md Step 11)',          '☐'),
        ('10', 'SMTP email notifications configured',         'Settings → Email (SMTP)',                   '☐'),
        ('11', 'SLA policies set per severity',               'Settings → SLA & Risk Appetite',           '☐'),
        ('12', 'Risk appetite documented',                    'Settings → SLA & Risk Appetite',           '☐'),
    ],
    col_widths=[0.8, 6.5, 5.5, 1.0],
)

# ── FINAL PAGE — FOOTER NOTE ─────────────────────────────────────────────────
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run('SecureOps Platform')
r.bold = True; r.font.size = Pt(16); r.font.color.rgb = ACCENT_BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('ISO/IEC 42001:2023 Implementation Guide')
r.font.size = Pt(12); r.font.color.rgb = MID_GRAY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f'Generated: {datetime.date.today().strftime("%d %B %Y")}')
r.font.size = Pt(10); r.font.color.rgb = MID_GRAY

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Use this platform responsibly — only on AI systems and networks you own or are authorised to govern.')
r.italic = True; r.font.size = Pt(9); r.font.color.rgb = MID_GRAY

# ── Save ─────────────────────────────────────────────────────────────────────
output = r'C:\Users\igigilashvili\Desktop\files\infosec-platform\ISO42001_SecureOps_Implementation_Guide.docx'
doc.save(output)
print(f'Saved: {output}')
